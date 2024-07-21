import io
import fitz  # PyMuPDF
from transformers import pipeline
from langdetect import detect
from docx import Document
import pandas as pd
import matplotlib.pyplot as plt
import streamlit as st

# Initialize translation pipeline with a different model
try:
    translator = pipeline("translation", model="Helsinki-NLP/opus-mt-en-fr", device=-1)
    # st.write("Translation pipeline initialized successfully.")  # Comment out this line
except Exception as e:
    st.error(f"Error initializing translation pipeline: {e}")
    translator = None

# Define function to read PDF and extract text
def extract_text_from_pdf(uploaded_file):
    pdf_document = fitz.open(stream=uploaded_file.read(), filetype="pdf")
    text = ""
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        text += page.get_text()
    pdf_document.close()
    return text

# Define function to detect language and translate text if necessary
def detect_and_translate(text):
    lang = detect(text)
    if lang != "en" and translator:
        try:
            translated_text = translator(text)[0]['translation_text']
            return translated_text, lang
        except Exception as e:
            st.error(f"Error during translation: {e}")
            return text, lang
    return text, lang

# Define function to clean up text
def clean_text(text):
    text = text.replace(' ', '\n')
    text = text.replace('\n\n', '\n')
    text = text.strip()
    return text

# Define function to extract skills from resume text
def extract_skills(text, skills):
    # Convert text to lowercase for case-insensitive matching
    text_lower = text.lower()
    skill_counts = {skill: text_lower.count(skill.lower()) for skill in skills}
    return skill_counts

# Define function to process resumes and extract skills
def process_resumes_and_extract_skills(uploaded_files, skills):
    resume_skills = {}
    
    for uploaded_file in uploaded_files:
        text = extract_text_from_pdf(uploaded_file)
        translated_text, lang = detect_and_translate(text)
        cleaned_text = clean_text(translated_text)
        skills_count = extract_skills(cleaned_text, skills)
        resume_skills[uploaded_file.name] = skills_count
        
    return resume_skills

# Streamlit app
st.title("Resume Analyzer")

# Add text input for required skills
required_skills_input = st.text_area("Enter required skill sets (comma-separated)", value="Python, Machine Learning, Data Analysis, SQL, Tableau")
required_skills = [skill.strip().lower() for skill in required_skills_input.split(',')]

uploaded_files = st.file_uploader("Choose PDF files", type="pdf", accept_multiple_files=True)
if uploaded_files:
    # Process resumes
    resume_skills = process_resumes_and_extract_skills(uploaded_files, required_skills)
    
    # Convert skills data to DataFrame
    skill_names = set()
    for skills in resume_skills.values():
        skill_names.update(skills.keys())
    
    skill_names = list(skill_names)
    data = []
    for resume, skills in resume_skills.items():
        row = {skill: skills.get(skill, 0) for skill in skill_names}
        row["Resume"] = resume
        data.append(row)
    
    df = pd.DataFrame(data)
    
    # Filter resumes based on required skills
    filtered_resumes = []
    for resume, skills in resume_skills.items():
        if any(skills.get(skill, 0) > 0 for skill in required_skills):
            filtered_resumes.append((resume, skills))
    
    # Convert filtered resumes data to DataFrame
    filtered_data = []
    for resume, skills in filtered_resumes:
        row = {skill: skills.get(skill, 0) for skill in skill_names}
        row["Resume"] = resume
        filtered_data.append(row)
    
    filtered_df = pd.DataFrame(filtered_data)
    
    # Plot the skills comparison
    st.write("Skills Comparison:")
    fig, ax = plt.subplots(figsize=(12, 8))
    if not filtered_df.empty:
        filtered_df.set_index('Resume').plot(kind='bar', ax=ax)
        ax.set_ylabel('Skill Count')
        ax.set_title('Skills Comparison Across Resumes')
        plt.xticks(rotation=45)
    else:
        st.write("No resumes match the required skills.")
    st.pyplot(fig)
    
    # Save results to a Word document and provide a download link
    if st.button("Save Results to Word Document"):
        doc = Document()
        doc.add_heading('Skills Comparison Report', 0)
        for resume, skills in filtered_resumes:
            doc.add_paragraph(f"Resume: {resume}")
            doc.add_paragraph(f"Skills:")
            for skill, count in skills.items():
                doc.add_paragraph(f"{skill}: {count}")
            doc.add_paragraph()
        doc_path = "skills_comparison_report.docx"
        doc.save(doc_path)
        
        with open(doc_path, "rb") as f:
            st.download_button(
                label="Download Skills Comparison Report",
                data=f,
                file_name=doc_path,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# Footer
st.markdown(
    """
    <style>
    .footer {
        position: fixed;
        bottom: 0;
        left: 0;
        width: 100%;
        background-color: #000000; /* Black background */
        color: #FFFFFF; /* White text */
        text-align: center;
        padding: 10px;
        font-size: 14px;
    }
    </style>
    <div class="footer">
        <p>Created by Sayan Sahu &copy; 2024</p>
    </div>
    """,
    unsafe_allow_html=True
)
